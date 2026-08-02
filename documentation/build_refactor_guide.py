from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "output" / "Capstone_Refactor_Implementation_Guide.docx"

BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
INK = "0B2545"
MUTED = "667085"
HEADER_FILL = "E8EEF5"
CALLOUT_FILL = "F4F6F9"
WARNING_FILL = "FFF7E6"


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    margins = tc_pr.first_child_found_in("w:tcMar")
    if margins is None:
        margins = OxmlElement("w:tcMar")
        tc_pr.append(margins)
    for side, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = margins.find(qn(f"w:{side}"))
        if node is None:
            node = OxmlElement(f"w:{side}")
            margins.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_cell_width(cell, dxa):
    tc_pr = cell._tc.get_or_add_tcPr()
    width = tc_pr.find(qn("w:tcW"))
    if width is None:
        width = OxmlElement("w:tcW")
        tc_pr.append(width)
    width.set(qn("w:w"), str(dxa))
    width.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths):
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.first_child_found_in("w:tblW")
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), "9360")
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.first_child_found_in("w:tblInd")
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")
    grid = table._tbl.tblGrid
    for col, width in zip(grid.gridCol_lst, widths):
        col.set(qn("w:w"), str(width))
    for row in table.rows:
        for cell, width in zip(row.cells, widths):
            set_cell_width(cell, width)
            set_cell_margins(cell)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def set_font(run, size=11, color=INK, bold=False, italic=False):
    run.font.name = "Calibri"
    run._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    run._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor.from_string(color)
    run.bold = bold
    run.italic = italic


def set_style(style, size, color, before, after, bold=False):
    style.font.name = "Calibri"
    style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    style.font.size = Pt(size)
    style.font.color.rgb = RGBColor.from_string(color)
    style.font.bold = bold
    pf = style.paragraph_format
    pf.space_before = Pt(before)
    pf.space_after = Pt(after)
    pf.line_spacing = 1.25


def add_text(doc, text, style=None, bold_prefix=None):
    p = doc.add_paragraph(style=style)
    if bold_prefix and text.startswith(bold_prefix):
        set_font(p.add_run(bold_prefix), bold=True)
        set_font(p.add_run(text[len(bold_prefix):]))
    else:
        set_font(p.add_run(text))
    return p


def add_bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.15
    set_font(p.add_run(text))
    return p


def add_number(doc, text):
    p = doc.add_paragraph(style="List Number")
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.15
    set_font(p.add_run(text))
    return p


def add_callout(doc, label, text, fill=CALLOUT_FILL):
    table = doc.add_table(rows=1, cols=1)
    set_table_geometry(table, [9360])
    cell = table.cell(0, 0)
    set_cell_shading(cell, fill)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    set_font(p.add_run(label + " "), color=DARK_BLUE, bold=True)
    set_font(p.add_run(text))
    doc.add_paragraph().paragraph_format.space_after = Pt(1)


def add_table(doc, headers, rows, widths):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    set_table_geometry(table, widths)
    for cell, text in zip(table.rows[0].cells, headers):
        set_cell_shading(cell, HEADER_FILL)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_font(p.add_run(text), size=9, color=DARK_BLUE, bold=True)
    for values in rows:
        cells = table.add_row().cells
        for index, (cell, text) in enumerate(zip(cells, values)):
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER if index == 0 else WD_ALIGN_PARAGRAPH.LEFT
            set_font(p.add_run(str(text)), size=8.5)
    doc.add_paragraph().paragraph_format.space_after = Pt(3)
    return table


def add_page_number(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run("Page ")
    set_font(run, size=8, color=MUTED)
    fld = OxmlElement("w:fldSimple")
    fld.set(qn("w:instr"), "PAGE")
    paragraph._p.append(fld)


def configure(doc):
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = doc.styles
    set_style(styles["Normal"], 11, INK, 0, 6)
    set_style(styles["Heading 1"], 16, BLUE, 14, 6, bold=True)
    set_style(styles["Heading 2"], 13, BLUE, 10, 4, bold=True)
    set_style(styles["Heading 3"], 12, DARK_BLUE, 8, 3, bold=True)
    if "Subtitle" not in styles:
        styles.add_style("Subtitle", WD_STYLE_TYPE.PARAGRAPH)
    set_style(styles["Subtitle"], 11, MUTED, 0, 10)

    header = section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.LEFT
    set_font(header.add_run("PAPER VENDO MACHINE | REFACTOR IMPLEMENTATION GUIDE"), size=8, color=MUTED, bold=True)
    footer = section.footer.paragraphs[0]
    add_page_number(footer)


def main():
    doc = Document()
    configure(doc)

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(3)
    set_font(p.add_run("IMPLEMENTATION GUIDE"), size=10, color=BLUE, bold=True)
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    set_font(p.add_run("Paper Vendo Machine Capstone Refactor"), size=23, color=INK, bold=True)
    p = doc.add_paragraph(style="Subtitle")
    set_font(p.add_run("Machine flow, database contract, web reporting, wiring additions, and deployment checks"), size=12, color=MUTED)

    add_table(doc, ["Prepared", "Scope", "Status"], [["2 August 2026", "Mega, ESP32, Supabase, PERN website", "Implementation refactor; hardware validation required"]], [1800, 4500, 3060])
    add_callout(doc, "Release rule.", "The machine must verify and record the full change payment before it receives a dispense plan. It never relies on a manual D button or an inferred website calculation.")

    doc.add_heading("What changed", level=1)
    add_bullet(doc, "A complete cart is now reserved atomically before a motor moves. The reservation includes products, user units, physical output quantity, stock, and exact change.")
    add_bullet(doc, "The Mega returns verified change first through a hopper sensor. Only CHANGE_OK lets the ESP32 request the dispense plan.")
    add_bullet(doc, "The transaction line stores units_requested separately from qty_requested and qty_dispensed. For paper, units are the customer choice and quantity is physical sheets.")
    add_bullet(doc, "The website reads stored units and price snapshots. It no longer reconstructs units from amount, sheet settings, or quantities after the sale.")
    add_bullet(doc, "The four paper channels are physical feeders by size. Budget and Standard are logical catalog products that may share a feeder only when their paper stock is physically identical.")

    doc.add_heading("Final transaction flow", level=1)
    steps = [
        "Customer inserts coins. The Mega maintains the credit in pesos and sends its exact cent value only when checkout begins.",
        "Customer selects Paper, then Budget or Standard, then one of four sizes; or selects one of three pen channels. The cart contains requested units, not calculated sheet counts.",
        "Mega inhibits the coin acceptor and sends one RESERVE message for the whole cart. ESP32 calls machine_reserve_transaction, which validates price, units, stock, shared feeder capacity, and exact change in one database transaction.",
        "ESP32 returns RESERVED with transaction ID, subtotal, and change due. Mega drives the change hopper and counts each exit-sensor edge.",
        "Only if every change coin is sensor-confirmed, Mega sends CHANGE_OK. ESP32 changes the transaction to CHANGE_PAID and returns the physical PLAN.",
        "Mega dispenses each plan item. Paper requires one exit-sensor cycle for each physical sheet; pens require the existing IR drop confirmation. Mega sends actual quantities in FINISH.",
        "ESP32 calls machine_finish_transaction. The database commits inventory consumption, marks every line DISPENSED or FAILED, and exposes the stored result to the website.",
    ]
    for step in steps:
        add_number(doc, step)

    add_callout(doc, "Failure behavior.", "If exact change cannot be confirmed, Mega sends CHANGE_FAIL and the reservation is cancelled. If change was confirmed but a dispenser sensor fails, the transaction is recorded as FAILED_DISPENSE with actual output; this requires operator resolution because change has already been given.", WARNING_FILL)

    doc.add_heading("Data model and calculation rule", level=1)
    add_table(doc, ["Concept", "Stored field", "Meaning"], [
        ["Customer choice", "units_requested", "How many sellable units the user selected."],
        ["Paper physical output", "qty_requested", "units_requested x sheets_per_unit_snapshot."],
        ["Confirmed output", "qty_dispensed", "Sensor-confirmed sheets or pens actually released."],
        ["Price evidence", "unit_price_cents", "Price captured at reservation; later setting edits cannot rewrite history."],
        ["Money", "credit/subtotal/change_*_cents", "Integer cent values avoid floating-point money errors."],
    ], [1800, 2520, 5040])
    add_text(doc, "Formula: subtotal_cents = sum(unit_price_cents x units_requested). change_due_cents = credit_received_cents - subtotal_cents. For a paper line, qty_requested = units_requested x sheets_per_unit_snapshot.", bold_prefix="Formula: ")

    doc.add_heading("Wiring additions only", level=1)
    add_text(doc, "Existing pen steppers, pen IR sensors, TFT, coin acceptor, OLED, servos, and Mega-ESP UART stay connected as they are. Add the following hardware paths:")
    add_table(doc, ["Hardware", "Mega pins", "Purpose"], [
        ["Paper drivers 1 to 4", "STEP/DIR: 32/33, 34/35, 36/37, 38/39; shared EN: 40", "Drive one adjustable paper feeder per physical size."],
        ["Paper exit sensors 1 to 4", "41, 42, 43, 44", "Confirm every sheet after it leaves its channel."],
        ["Change hopper motor driver", "14", "Releases reserved P1 change coins."],
        ["Change hopper coin sensor", "15", "Confirms every released change coin before dispensing."],
        ["Coin acceptor inhibit", "16", "Freezes coin acceptance after checkout starts so the credit snapshot cannot change mid-transaction."],
        ["Mega to ESP UART", "Mega Serial1 18/19 to ESP Serial2 16/17 through level shifting", "Carries the reservation and result protocol; retain existing UART path but protect ESP 3.3 V RX."],
    ], [2160, 3240, 3960])
    add_callout(doc, "Mechanical constraint.", "Four paper compartments support eight catalog choices only when Budget and Standard for the same size use the same physical paper. If brands differ in paper stock, color, weight, or finish, use eight separate physical feeders and map each product to its own channel.", WARNING_FILL)

    doc.add_page_break()
    doc.add_heading("File-by-file implementation summary", level=1)
    add_text(doc, "Line ranges below refer to the refactored files in this workspace. They identify the main implementation areas, not every supporting line.")
    add_table(doc, ["File", "Key lines", "Specific change and purpose"], [
        ["Cloud_Paper_Vendo.sql", "1-328; 152-326", "Clean development schema: physical channels, exact-change inventory, transaction header and lines, atomic reserve/change/finish RPCs. Purpose: authoritative, auditable machine data flow."],
        ["updated_arduino_mega.ino", "20-36; 178-226; 571-693; 949-1258", "Added paper/hopper/coin-inhibit pins, brand selection, matching local display prices, whole-cart reservation, verified change, physical paper and pen sensor routines, and FINISH results. Purpose: stable credit snapshot, change before dispense, and sensor-confirmed machine flow."],
        ["production_esp32.ino", "1-229; 125-212", "New gateway protocol and Supabase RPC calls. Purpose: translate serial cart/result messages into database state changes."],
        ["backend/routes/machine.js", "1-276; 102-257", "Flattens stored transaction lines, returns units and physical output, updates physical inventories, and calculates analytics from completed stored units. Purpose: stop frontend inference."],
        ["frontend/src/pages/Transactions.jsx", "1-72; 59-60", "Displays stored units, expected/actual output, snapshots, and line status. Purpose: transaction auditability."],
        ["frontend/src/pages/Reports.jsx", "56-145; 319-340", "Uses units_requested and stored product snapshots in reports/exports. Purpose: correct historical reports after settings change."],
        ["backend/routes/auth.js", "30-42", "Removes plaintext password fallback. Purpose: database seed and login use bcrypt only."],
    ], [2250, 1350, 5760])

    doc.add_heading("Deployment and calibration checklist", level=1)
    checks = [
        "Back up any real Supabase data. Cloud_Paper_Vendo.sql is a clean-reset development script and drops the listed tables.",
        "Run the SQL on the target development database, then enter real stock, prices, sheets per unit, and change-hopper coin count.",
        "Place the real Wi-Fi and Supabase values in production_esp32.ino. Do not place credentials in the guide or source control.",
        "Install four paper modules and calibrate each driver direction, step timing, separator, and exit sensor. A paper sheet counts only on clear -> blocked -> clear.",
        "Install a motor driver between Mega D14 and the hopper motor; never drive the motor from a Mega pin. Adjust active level if the module is active-low.",
        "Test each exact-change case with empty, low, jammed, and normal hopper states before customer testing.",
        "Run full-path tests: one pen, one paper unit with multiple sheets, mixed cart, zero change, valid change, no exact change, paper sensor failure, pen IR failure, Wi-Fi loss before reserve, and Wi-Fi loss after change payment.",
    ]
    for check in checks:
        add_bullet(doc, check)

    doc.add_heading("Verification completed in this refactor", level=1)
    add_bullet(doc, "Backend JavaScript syntax checked with Node.")
    add_bullet(doc, "Frontend production build completed successfully with Vite.")
    add_bullet(doc, "Mega sketch structure checked for balanced braces. Hardware compilation and physical tests remain required because this workspace does not include Arduino CLI/hardware access.")
    add_bullet(doc, "DOCX and PDF are rendered and visually inspected before delivery.")

    doc.add_heading("Operator acceptance criteria", level=1)
    add_bullet(doc, "No good is released unless the database reservation succeeds and the change sensor has verified the change due.")
    add_bullet(doc, "A website transaction shows the user-selected units, actual physical output, price snapshot, status, and timestamp without recomputing them.")
    add_bullet(doc, "Paper inventory reduces by confirmed physical sheets; pen inventory reduces by confirmed pens.")
    add_bullet(doc, "A sensor-short dispense is visible as a failed line and failed transaction, rather than silently recorded as success.")

    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    main()
